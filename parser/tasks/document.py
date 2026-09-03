import logging
import os
from chunk_reindex import reindex_chunks_preserving_links
from connectors.folder import extract_pdf_pages
from core import config
from db import SessionLocal
from models.database import KnowledgeSource, DocumentChunk
from code_parser import CodeParser
from ollama_client import get_embedding, ensure_model_pulled

logger = logging.getLogger(__name__)

async def process_local_document_async(source_id: int, file_path: str):
    """
    Parses and indexes a locally uploaded file (PDF, Word, or plain text).
    
    1. Extracts text content depending on the file format (PDF parsing via the
       shared `connectors.folder.extract_pdf_pages` incl. OCR fallback for
       image-only PDFs, Word parsing with python-docx, or default raw text
       reading).
    2. Removes invalid characters (null bytes).
    3. Triggers embedding model pulling in Ollama.
    4. Chunks the document content using CodeParser.
    5. Clears previous chunks for this specific knowledge source.
    6. Generates vector embeddings for each chunk and saves them.
    """
    db = SessionLocal()
    source = db.query(KnowledgeSource).filter(KnowledgeSource.id == source_id).first()
    if not source:
        logger.warning(f"Local document source {source_id} not found.")
        return

    if source.sync_status == "cancelled":
        logger.info(f"Local document source {source_id} was cancelled before processing.")
        db.close()
        return

    if source.sync_status == "syncing":
        logger.info(f"Verarbeitung für Quelle {source_id} läuft bereits, überspringe.")
        return

    from datetime import datetime, timezone
    def log_event(message: str):
        timestamp = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')
        full_msg = f"[{timestamp}] {message}\n"
        logger.info(message)
        source.sync_log = (source.sync_log or "") + full_msg
        db.commit()

    try:
        source.sync_status = "syncing"
        source.last_error = None
        source.sync_log = ""
        db.commit()

        log_event(f"Starte Verarbeitung der lokalen Datei '{file_path}' (ID: {source_id})...")
        
        # 1. Read and extract content depending on file type.
        # `pages` holds (page_number, text) tuples -- page_number is the 1-based
        # PDF page for formats that have one, None otherwise. Chunking per page
        # (instead of joining all pages into one blob first) keeps start_line/
        # end_line meaningful relative to an actual page, and lets us tag each
        # chunk with the page it came from -- needed to answer "was steht auf
        # Seite X" questions, which is otherwise structurally impossible once
        # pages are flattened into a single string.
        ext = os.path.splitext(file_path)[1].lower()
        pages = []

        if ext == ".pdf":
            try:
                log_event("Lese PDF-Dokument ein...")
                pages = [(page_no, text) for page_no, text in extract_pdf_pages(file_path) if text]
            except Exception as e:
                log_event(f"Fehler beim Lesen der PDF-Datei, versuche Plaintext-Fallback: {e}")
                try:
                    with open(file_path, "r", errors="ignore") as f:
                        fallback = f.read()
                    if fallback:
                        pages.append((None, fallback))
                except Exception:
                    pass
        elif ext in [".docx", ".doc"]:
            try:
                log_event("Lese Word-Dokument (.docx) ein...")
                import docx
                doc = docx.Document(file_path)
                text_parts = []
                for para in doc.paragraphs:
                    text_parts.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        text_parts.append(" | ".join(row_text))
                pages.append((None, "\n".join(text_parts)))
            except Exception as e:
                log_event(f"Fehler beim Lesen des Word-Dokuments: {e}")
        else:
            try:
                log_event("Lese Textdatei ein...")
                with open(file_path, "r", errors="ignore") as f:
                    pages.append((None, f.read()))
            except Exception as e:
                log_event(f"Fehler beim Lesen der Datei: {e}")

        # Clean null bytes
        pages = [(page_no, text.replace("\x00", "")) for page_no, text in pages]
        total_chars = sum(len(text) for _, text in pages)
        if not any(text.strip() for _, text in pages):
            log_event(f"Kein Textinhalt aus {file_path} extrahiert. Abort.")
            raise Exception("Kein Textinhalt extrahiert.")

        log_event(f"{total_chars} Zeichen Text erfolgreich extrahiert ({len(pages)} Seite(n)).")

        # 2. Ensure embedding model is pulled
        log_event(f"Stelle sicher, dass das Einbettungs-Modell '{config.EMBED_MODEL}' bereit ist...")
        await ensure_model_pulled(config.EMBED_MODEL)

        # 3. Chunk and embed content, page by page so a chunk never blends text
        # from two different PDF pages together.
        lang = "markdown" if ext == ".md" else "text"
        parser = CodeParser(lang)
        chunks = []
        for page_no, text in pages:
            for chunk in parser.chunk_file(text):
                chunk["page"] = page_no
                chunks.append(chunk)
        log_event(f"Datei in {len(chunks)} Chunks aufgeteilt. Starte Einbettung...")

        def build_chunk(chunk, embedding):
            return DocumentChunk(
                project_id=source.project_id,
                source_id=source.id,
                file_path=source.name,
                content=chunk["content"],
                start_line=chunk["start_line"],
                end_line=chunk["end_line"],
                embedding=embedding,
                metadata_json={"language": lang, "page": chunk.get("page")}
            )

        async def embed_content(content):
            return await get_embedding(content, model=config.EMBED_MODEL)

        def on_embed_error(chunk, e):
            log_event(f"Fehler beim Erzeugen des Vektors für Chunk: {e}")

        embedded_chunks_count = await reindex_chunks_preserving_links(
            db,
            source_id=source_id,
            file_path=source.name,
            chunks=chunks,
            build_chunk=build_chunk,
            embed_content=embed_content,
            on_embed_error=on_embed_error,
        )

        db.commit()
        source.sync_status = "completed"
        source.last_synced_at = datetime.now(timezone.utc)
        db.commit()
        log_event(f"Datei '{file_path}' erfolgreich indiziert ({embedded_chunks_count} Vektor-Chuncks erzeugt).")

    except Exception as e:
        error_msg = str(e)
        log_event(f"Kritischer Fehler bei Dateiverarbeitung: {error_msg}")
        source.sync_status = "error"
        source.last_error = error_msg
        db.commit()
    finally:
        db.close()
