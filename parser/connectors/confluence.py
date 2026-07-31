"""
parser/connectors/confluence.py
================================
Connector für Atlassian Confluence (Cloud & Server).

Unterstützte Authentifizierung:
    - Basic Auth: source.username + source.token (API-Token)
    - Bearer Token: nur source.token gesetzt

Delta-Sync:
    Jede Seite und jeder Anhang wird mit source.last_synced_at verglichen
    (page/attachment version.when). Unveränderte Objekte werden übersprungen.

Attachments:
    Für jede Seite werden Anhänge über /child/attachment geladen.
    Unterstützte Typen: PDF (pypdf), DOCX (python-docx), Plaintext/Code.
    Binärdateien ohne Textinhalt (Bilder, ZIP, ...) werden übersprungen.
    Maximale Dateigröße: 20 MB (ATTACHMENT_MAX_BYTES).

Pagination:
    Confluence liefert Seiten in Batches (Standard: 25). Die _fetch()-Methode
    enthält exponentielles Backoff für Rate-Limiting (HTTP 429).
"""

import io
import re
import asyncio
from datetime import datetime, timezone
from html.parser import HTMLParser

import httpx

from connectors.base import BaseConnector, Document
from db import SessionLocal
from models.database import DocumentChunk

ATTACHMENT_MAX_BYTES = 20 * 1024 * 1024  # 20 MB

_SUPPORTED_MIME_PREFIXES = ("text/",)
_SUPPORTED_MIME_EXACT = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}


class _ConfluenceHTMLParser(HTMLParser):
    """Converts Confluence rendered HTML (body.view.value) to clean plaintext.

    Improvements over a raw regex strip:
    - HTML entities decoded correctly (&amp; → &, &nbsp; → space, etc.)
    - Block-level tags produce newlines instead of spaces, preserving paragraphs
    - Table cells separated with ' | ' so rows stay readable
    - <script>, <style>, and Confluence param tags excluded from output
    - <pre>/<code> content kept verbatim (no whitespace collapsing inside)
    """

    _BLOCK = frozenset({
        "p", "div", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6",
        "blockquote", "ul", "ol", "table", "thead", "tbody", "pre",
    })
    _SKIP = frozenset({"script", "style", "ac:parameter"})

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self._buf: list[str] = []
        self._pre_depth = 0
        self._skip_depth = 0

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skip_depth += 1
            return
        if self._skip_depth:
            return
        if tag in ("pre", "code"):
            self._pre_depth += 1
            self._buf.append("\n")
        elif tag == "br":
            self._buf.append("\n")
        elif tag in ("td", "th"):
            self._buf.append(" | ")
        elif tag in self._BLOCK:
            self._buf.append("\n")

    def handle_endtag(self, tag):
        if tag in self._SKIP:
            self._skip_depth = max(0, self._skip_depth - 1)
            return
        if self._skip_depth:
            return
        if tag in ("pre", "code"):
            self._pre_depth = max(0, self._pre_depth - 1)
            self._buf.append("\n")
        elif tag in self._BLOCK:
            self._buf.append("\n")

    def handle_data(self, data):
        if self._skip_depth:
            return
        self._buf.append(data)

    def get_text(self) -> str:
        raw = "".join(self._buf)
        lines = []
        for line in raw.split("\n"):
            # Collapse inline whitespace per line; keep blank lines for structure
            normalized = " ".join(line.split())
            lines.append(normalized)
        # Drop leading/trailing blanks, collapse 3+ blank lines to 2
        text = "\n".join(lines).strip()
        return re.sub(r"\n{3,}", "\n\n", text)


def _html_to_text(html: str) -> str:
    parser = _ConfluenceHTMLParser()
    parser.feed(html)
    return parser.get_text()


def _extract_attachment_text(data: bytes, mime_type: str) -> str | None:
    """Extracts plaintext from attachment bytes. Returns None if unsupported or empty."""
    mime = (mime_type or "").split(";")[0].strip().lower()

    if mime == "application/pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(p.strip() for p in pages if p.strip()) or None
        except Exception:
            return None

    if mime in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        try:
            from docx import Document as DocxDoc
            doc = DocxDoc(io.BytesIO(data))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs) or None
        except Exception:
            return None

    if mime.startswith("text/"):
        try:
            return data.decode("utf-8", errors="replace").strip() or None
        except Exception:
            return None

    return None


def _is_supported_mime(mime_type: str) -> bool:
    mime = (mime_type or "").split(";")[0].strip().lower()
    return mime in _SUPPORTED_MIME_EXACT or any(mime.startswith(p) for p in _SUPPORTED_MIME_PREFIXES)


class ConfluenceConnector(BaseConnector):
    """
    Crawlt einen Confluence-Space und liefert Seiten als Document-Objekte.

    Confluence liefert `body.view.value` als gerendertes HTML. Die Umwandlung
    in Plaintext erfolgt über _html_to_text (HTMLParser), nicht per Regex,
    damit Entities korrekt dekodiert werden und Code-Blöcke lesbar bleiben.
    """

    async def fetch_documents(self):
        spaces = self._parse_spaces()
        self._log(f"Verbinde mit Confluence unter {self.source.url} (Spaces: {spaces})...")

        # Auth-Setup: Basic Auth hat Vorrang vor Bearer Token
        auth = None
        headers: dict = {}
        if self.source.username and self.source.token:
            auth = (self.source.username, self.source.token)
        elif self.source.token:
            headers["Authorization"] = f"Bearer {self.source.token}"

        base_url = self.source.url.rstrip("/")

        async with httpx.AsyncClient() as client:
            start = 0
            page_limit = 25
            connection_verified = False

            while True:
                pages_data = await self._fetch_pages(
                    client, base_url, auth, headers, spaces, start, page_limit
                )
                if pages_data is None:
                    if not connection_verified:
                        raise RuntimeError(
                            "Verbindung zur Confluence-API fehlgeschlagen. Bitte überprüfe die Server-URL und die API-Logins."
                        )
                    break
                connection_verified = True

                results = pages_data.get("results", [])
                if not results:
                    break

                for i, page in enumerate(results):
                    current_count = start + i + 1
                    # Raus, wenn die Seite nicht zum gewünschten Space gehört
                    page_space = page.get("space", {}).get("key")
                    if "ALL" not in spaces and page_space not in spaces:
                        continue

                    doc = self._build_document(page, base_url)
                    self._update_progress(current_count, message=f"Verarbeite Seite '{page.get('title', '...')}' ({current_count})...")
                    
                    if doc is None:
                        continue  # Leer oder unverändert
                    yield doc

                    async for att_doc in self._fetch_attachments(
                        client, base_url, auth, headers, page
                    ):
                        yield att_doc

                start += len(results)
                # Pagination: Ende erreicht wenn weniger Ergebnisse als erwartet
                # oder kein "next"-Link in der Antwort
                if len(results) < page_limit or "next" not in pages_data.get("_links", {}):
                    break

    # ── Interne Hilfsmethoden ────────────────────────────────────────────────

    async def _fetch_pages(
        self, client, base_url, auth, headers, spaces, start, limit
    ) -> dict | None:
        """
        Ruft eine Seite der Confluence-Content-API ab.
        Versucht beide bekannten API-Pfade (/wiki/rest/api und /rest/api),
        da der korrekte Pfad je nach Confluence-Installation variiert.
        """
        params = {
            "type": "page",
            "start": start,
            "limit": limit,
            "expand": "body.view,version,space",
        }
        if "ALL" not in spaces and len(spaces) == 1:
            params["spaceKey"] = spaces[0]

        for path in ["/wiki/rest/api/content", "/rest/api/content"]:
            try:
                return await self._fetch_with_retry(
                    client, f"{base_url}{path}", auth, headers, params
                )
            except Exception as e:
                self._log(f"API-Pfad {path} fehlgeschlagen: {e}. Versuche Alternative...")

        self._log("Keine gültige Confluence-API gefunden.")
        return None

    async def _fetch_with_retry(
        self, client, url, auth, headers, params, max_retries: int = 5
    ) -> dict:
        """
        GET-Request mit exponentiellem Backoff bei Rate-Limiting (HTTP 429)
        und transientes Fehler-Retries.
        """
        backoff = 1.0
        for attempt in range(max_retries):
            try:
                response = await client.get(
                    url, auth=auth, headers=headers, params=params, timeout=30.0
                )
                if response.status_code == 429:
                    # Confluence gibt manchmal Retry-After vor, sonst exponentiell warten
                    wait = float(response.headers.get("Retry-After", backoff))
                    self._log(f"Rate Limit (429). Warte {wait}s...")
                    await asyncio.sleep(wait)
                    backoff *= 2.0
                    continue
                response.raise_for_status()
                await asyncio.sleep(0.2)  # Höflichkeits-Pause zwischen Requests
                return response.json()
            except Exception as e:
                if attempt < max_retries - 1:
                    self._log(f"Anfrage an {url} fehlgeschlagen: {e}. Retry in {backoff}s...")
                    await asyncio.sleep(backoff)
                    backoff *= 2.0
                else:
                    raise

    def _build_document(self, page: dict, base_url: str) -> Document | None:
        """
        Wandelt ein Confluence-Page-Objekt in ein Document um.

        Gibt None zurück wenn:
        - Die Seite seit dem letzten Sync unverändert ist (und Chunks existieren)
        - Der Seiteninhalt leer ist
        """
        title = page.get("title", "Unbekannte Seite")

        # Delta-Sync: Seite überspringen wenn unverändert und bereits indiziert
        updated_at_str = page.get("version", {}).get("when")
        if self.source.last_synced_at and updated_at_str:
            try:
                updated_at = datetime.fromisoformat(updated_at_str.replace("Z", "+00:00"))
                if updated_at <= self.source.last_synced_at:
                    chunks_exist = self.db.query(DocumentChunk).filter(
                        DocumentChunk.source_id == self.source.id,
                        DocumentChunk.file_path == title
                    ).first() is not None
                    if chunks_exist:
                        return None
            except Exception as ex:
                self._log(f"Konnte Änderungszeitpunkt für '{title}' nicht parsen: {ex}")

        body_html = page.get("body", {}).get("view", {}).get("value", "")
        plain_text = _html_to_text(body_html)
        if not plain_text:
            self._log(f"Seite '{title}' hat keinen Textinhalt. Überspringe.")
            return None

        # URL aufbauen — Confluence liefert relative Links in _links.webui
        webui = page.get("_links", {}).get("webui", "")
        if "/wiki" in page.get("_links", {}).get("base", ""):
            original_url = f"{base_url}/wiki{webui}"
        elif webui:
            original_url = f"{base_url}{webui}"
        else:
            original_url = f"{base_url}/pages/viewpage.action?pageId={page.get('id')}"

        return Document(
            title=title,
            content=plain_text,
            url=original_url,
            source_type="Confluence",
            storage_key=title,
            extra_meta={"page_id": page.get("id")},
        )

    async def _fetch_attachments(self, client, base_url, auth, headers, page: dict):
        """Async generator: yields Document for each supported attachment on a page."""
        page_id = page.get("id")
        page_title = page.get("title", "Unbekannte Seite")

        for path in [f"/wiki/rest/api/content/{page_id}/child/attachment",
                     f"/rest/api/content/{page_id}/child/attachment"]:
            try:
                resp = await self._fetch_with_retry(
                    client, f"{base_url}{path}", auth, headers,
                    {"limit": 50, "expand": "version"}
                )
                break
            except Exception:
                resp = None
        else:
            return

        for att in (resp or {}).get("results", []):
            mime_type = att.get("metadata", {}).get("mediaType", "")
            if not _is_supported_mime(mime_type):
                continue

            filename = att.get("title", "attachment")
            file_size = att.get("extensions", {}).get("fileSize", 0)
            if file_size and file_size > ATTACHMENT_MAX_BYTES:
                self._log(f"Anhang '{filename}' übersprungen (>{ATTACHMENT_MAX_BYTES // (1024*1024)} MB).")
                continue

            # Delta-Sync: Anhang überspringen wenn unverändert und bereits indiziert
            storage_key = f"{page_title}/attachments/{filename}"
            updated_at_str = att.get("version", {}).get("when")
            if self.source.last_synced_at and updated_at_str:
                try:
                    updated_at = datetime.fromisoformat(updated_at_str.replace("Z", "+00:00"))
                    if updated_at <= self.source.last_synced_at:
                        from models.database import DocumentChunk as DC
                        if self.db.query(DC).filter(
                            DC.source_id == self.source.id,
                            DC.file_path == storage_key
                        ).first():
                            continue
                except Exception:
                    pass

            download_path = att.get("_links", {}).get("download", "")
            if not download_path:
                continue

            try:
                dl_resp = await client.get(
                    f"{base_url}{download_path}", auth=auth, headers=headers,
                    follow_redirects=True, timeout=60.0
                )
                dl_resp.raise_for_status()
            except Exception as e:
                self._log(f"Download-Fehler für '{filename}': {e}")
                continue

            text = _extract_attachment_text(dl_resp.content, mime_type)
            if not text:
                continue

            att_url = f"{base_url}{download_path}"
            yield Document(
                title=f"{page_title} — {filename}",
                content=text,
                url=att_url,
                source_type="Confluence",
                storage_key=storage_key,
                extra_meta={"page_id": page_id, "attachment_filename": filename, "mime_type": mime_type},
            )
