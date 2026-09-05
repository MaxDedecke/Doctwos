"""
backend/api/knowledge_sources.py
==================================
Verwaltung von Wissensquellen (Git, Confluence, Jira, WebDAV, FolderWatch, Upload).

Endpoint-Übersicht:
    POST   /knowledge-sources            — Quelle anlegen + Sync triggern
    GET    /knowledge-sources            — Alle Quellen auflisten
    DELETE /knowledge-sources/{id}       — Quelle + Chunks löschen
    POST   /knowledge-sources/{id}/sync  — Manuellen Re-Sync starten
    GET    /knowledge-sources/{id}/resolve?url=… — URL → gerendertes HTML
    GET    /knowledge-sources/{id}/content — Dateiinhalt (PDF/Word/Text) extrahieren
    GET    /knowledge-sources/{id}/raw    — Datei als Binary ausliefern
    POST   /knowledge-sources/upload      — Lokale Datei hochladen + parsen

Celery-Tasks:
    "process_knowledge_source" — Universeller Task für Git/Confluence/Jira/WebDAV
    "process_local_document"   — Für lokale Datei-Uploads

Render-Logik (/resolve):
    Der Endpoint lädt den Quell-Inhalt live von der jeweiligen API
    und gibt ihn als styled HTML zurück. Das erlaubt der UI, externe Seiten
    inline anzuzeigen ohne einen Browser-Tab zu öffnen.
"""

import logging
import os
import re
import shutil
from typing import Optional

import httpx
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse
from sqlalchemy import or_
from sqlalchemy.orm import Session

from api.schemas import FolderWatchCreate, KnowledgeSourceCreate, KnowledgeSourceUpdate, GitSourceCreate
from api.serializers import serialize_source
from core.config import celery_app, UPLOADS_DIR, REPOS_ROOT
from core.db_setup import get_db
from models.database import DocumentChunk, JobCenterDismissal, KnowledgeSource, Project, Team, User
from core.auth_dependency import get_current_user
from core.teams import get_visible_team_ids, assert_team_visible, is_admin, require_admin, DEFAULT_TEAM_NAME
from core.projects import assert_knowledge_source_visible, assert_project_visible, get_visible_project_ids
from core.tracing import get_trace_id
from services.job_control import send_tracked_task

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/knowledge-sources", tags=["knowledge_sources"])


MAX_KNOWLEDGE_SOURCES_PER_REPO = 2

# Erlaubte Auto-Sync-Intervalle (Minuten), passend zu den UI-Optionen
# 1h / 6h / 24h / manuell. 0 = nur manuell (kein automatischer Sync).
ALLOWED_SYNC_INTERVALS = {0, 60, 360, 1440}
# Default für neu angelegte Quellen: nur manuell. Auto-Sync gegen externe
# CDEs/Wikis ohne explizites Opt-in würde sonst unbemerkt wiederkehrend
# Requests gegen Kunden-APIs auslösen (Rate-Limits, Kosten) und Re-Index-Läufe
# anstoßen, bevor der Nutzer das erste Ergebnis überhaupt gesehen hat.
DEFAULT_SYNC_INTERVAL = 0

# Harte Zeichengrenze pro Notiz — die Notiz landet bei jeder Chat-Anfrage
# ungekürzt im System-Prompt (siehe services/source_context.py), muss also
# bewusst kurz/kuratiert bleiben statt zum Ablageort für ganze Glossare zu
# werden (dafür sind reguläre Wissensquellen mit RAG-Retrieval da).
SOURCE_CONTEXT_NOTE_MAX_CHARS = 2000


def _validate_sync_interval(value: Optional[int]) -> int:
    """Gibt ein gültiges Intervall zurück; None → Default. Wirft 400 bei ungültigem Wert."""
    if value is None:
        return DEFAULT_SYNC_INTERVAL
    if value not in ALLOWED_SYNC_INTERVALS:
        raise HTTPException(
            status_code=400,
            detail=f"Ungültiges Sync-Intervall {value}. Erlaubt: {sorted(ALLOWED_SYNC_INTERVALS)} (0 = nur manuell).",
        )
    return value


def _validate_context_note(value: Optional[str]) -> Optional[str]:
    """Trimmt die Notiz; eine leere Notiz wird als NULL gespeichert. Wirft 400 bei Überlänge."""
    if value is None:
        return None
    stripped = value.strip()
    if not stripped:
        return None
    if len(stripped) > SOURCE_CONTEXT_NOTE_MAX_CHARS:
        raise HTTPException(
            status_code=400,
            detail=f"Kontext-Notiz zu lang ({len(stripped)} Zeichen, erlaubt: {SOURCE_CONTEXT_NOTE_MAX_CHARS}).",
        )
    return stripped


def _check_knowledge_source_cap(project_id: Optional[int], db: Session):
    pass



def _resolve_team_id(
    project_id: Optional[int],
    db: Session,
    user: User,
    client_team_id: Optional[int] = None
) -> int:
    if project_id is not None:
        proj = db.query(Project).filter(Project.id == project_id).first()
        if not proj:
            raise HTTPException(status_code=404, detail="Projekt nicht gefunden")
        assert_team_visible(proj.team_id, user, db, "Projekt nicht gefunden")
        assert_project_visible(project_id, user, db)
        return proj.team_id

    team_ids = get_visible_team_ids(user, db)
    if client_team_id is not None:
        if team_ids is not None and client_team_id not in team_ids:
            raise HTTPException(status_code=403, detail="Kein Zugriff auf dieses Team")
        return client_team_id
    else:
        if is_admin(user):
            default_team = db.query(Team.id).filter(Team.name == DEFAULT_TEAM_NAME).scalar()
            if default_team:
                return default_team
            raise HTTPException(status_code=403, detail="Standard-Team nicht gefunden")
        else:
            if team_ids and len(team_ids) == 1:
                return team_ids[0]
            else:
                raise HTTPException(status_code=403, detail="Team-Auswahl erforderlich")


@router.post("")
def create_knowledge_source(
    source: KnowledgeSourceCreate,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    team_id = _resolve_team_id(source.project_id, db, user, source.team_id)
    _check_knowledge_source_cap(source.project_id, db)
    db_source = KnowledgeSource(
        name=source.name, type=source.type, url=source.url,
        username=source.username, token=source.token,
        project_id=source.project_id, spaces=source.spaces,
        sync_interval_minutes=_validate_sync_interval(source.sync_interval_minutes),
        context_note=_validate_context_note(source.context_note),
        team_id=team_id
    )
    db.add(db_source)
    db.commit()
    db.refresh(db_source)

    # Einheitlicher Task für alle Web-Connector-Typen (Git/Confluence/Jira/WebDAV)
    if db_source.type and db_source.type.lower() in ("confluence", "jira", "webdav", "git"):
        send_tracked_task(db, db_source, "process_knowledge_source", [db_source.id], {"trace_id": get_trace_id()})

    return serialize_source(db_source)


@router.get("")
def get_knowledge_sources(
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    team_ids = get_visible_team_ids(user, db)
    project_ids = get_visible_project_ids(user, db)
    q = db.query(KnowledgeSource)
    if team_ids is not None:
        q = q.filter(KnowledgeSource.team_id.in_(team_ids))
    if project_ids is not None:
        q = q.filter(or_(KnowledgeSource.project_id.in_(project_ids), KnowledgeSource.project_id.is_(None)))
    return [serialize_source(s) for s in q.all()]


@router.delete("/{source_id}")
def delete_knowledge_source(
    source_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    db_source = db.query(KnowledgeSource).filter(KnowledgeSource.id == source_id).first()
    if not db_source:
        raise HTTPException(status_code=404, detail="Wissensquelle nicht gefunden")
    assert_knowledge_source_visible(db_source, user, db)

    db.query(DocumentChunk).filter(DocumentChunk.source_id == source_id).delete()

    if db_source.type == "Local" and isinstance(db_source.spaces, dict):
        file_path = db_source.spaces.get("path")
        if file_path:
            full_path = os.path.join(UPLOADS_DIR, file_path)
            if os.path.exists(full_path):
                try:
                    os.remove(full_path)
                except Exception as e:
                    logger.error(f"Fehler beim Löschen von {full_path}: {e}")

    db.delete(db_source)
    db.commit()
    return {"message": "Wissensquelle erfolgreich gelöscht"}


@router.post("/{source_id}/sync")
def sync_knowledge_source(
    source_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    db_source = db.query(KnowledgeSource).filter(KnowledgeSource.id == source_id).first()
    if not db_source:
        raise HTTPException(status_code=404, detail="Wissensquelle nicht gefunden")
    assert_knowledge_source_visible(db_source, user, db)

    if db_source.type and db_source.type.lower() in ("confluence", "jira", "folderwatch", "webdav", "git"):
        send_tracked_task(db, db_source, "process_knowledge_source", [db_source.id], {"trace_id": get_trace_id()})
    else:
        raise HTTPException(status_code=400, detail="Synchronisierung wird für diesen Quelltyp nicht unterstützt")

    return {"message": "Synchronisierung der Wissensquelle gestartet", "source_id": source_id}


@router.post("/{source_id}/reindex")
def reindex_knowledge_source(
    source_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_admin),
):
    """Queue an admin-only full reindex without recreating the source."""
    db_source = db.query(KnowledgeSource).filter(KnowledgeSource.id == source_id).first()
    if not db_source:
        raise HTTPException(status_code=404, detail="Wissensquelle nicht gefunden")
    if (db_source.type or "").lower() != "git":
        raise HTTPException(status_code=400, detail="Vollständige Neu-Analyse wird nur für Git unterstützt")
    if db_source.sync_status in {"pending", "syncing"}:
        raise HTTPException(status_code=409, detail="Für diese Wissensquelle läuft bereits eine Analyse")

    db_source.sync_status = "pending"
    db_source.progress = 0
    db_source.parsed_files = 0
    db_source.total_files = 0
    db_source.progress_message = "Vollständige Neu-Analyse in Warteschlange…"
    db_source.last_error = None
    db_source.sync_log = ""
    db.query(JobCenterDismissal).filter(
        JobCenterDismissal.kind == "source",
        JobCenterDismissal.job_id == db_source.id,
    ).delete(synchronize_session=False)
    db.commit()
    send_tracked_task(
        db, db_source, "process_knowledge_source", [db_source.id],
        {"force_reindex": True, "trace_id": get_trace_id()},
    )
    return {"message": "Vollständige Neu-Analyse gestartet", "source_id": source_id}


@router.patch("/{source_id}")
def update_knowledge_source(
    source_id: int,
    update: KnowledgeSourceUpdate,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Aktualisiert Auto-Sync-Intervall und/oder Kontext-Notiz einer Wissensquelle.

    Nur Felder aktualisieren, die der Client tatsächlich mitgeschickt hat
    (exclude_unset) — sonst würde z.B. ein reines Notiz-Update das nicht
    mitgesendete sync_interval_minutes über _validate_sync_interval(None)
    unbeabsichtigt auf den Default (nur manuell) zurücksetzen.
    """
    db_source = db.query(KnowledgeSource).filter(KnowledgeSource.id == source_id).first()
    if not db_source:
        raise HTTPException(status_code=404, detail="Wissensquelle nicht gefunden")
    assert_knowledge_source_visible(db_source, user, db)

    fields = update.model_dump(exclude_unset=True)
    if "sync_interval_minutes" in fields:
        db_source.sync_interval_minutes = _validate_sync_interval(fields["sync_interval_minutes"])
    if "context_note" in fields:
        db_source.context_note = _validate_context_note(fields["context_note"])
    db.commit()
    db.refresh(db_source)
    return serialize_source(db_source)


@router.get("/{source_id}/resolve")
def resolve_knowledge_source_url(
    source_id: int,
    url: str,
    theme: str = "dark",
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """
    Löst eine (ggf. nicht-HTTP) URL zu gerendertem HTML auf.

    Wenn url kein http/https-Link ist, wird sie als Seitentitel interpretiert
    und über DocumentChunk.metadata_json.url in die echte URL übersetzt.
    """
    db_source = db.query(KnowledgeSource).filter(KnowledgeSource.id == source_id).first()
    if not db_source:
        raise HTTPException(status_code=404, detail="Wissensquelle nicht gefunden")
    assert_knowledge_source_visible(db_source, user, db)

    if url and not (url.startswith("http://") or url.startswith("https://")):
        chunk = db.query(DocumentChunk).filter(
            DocumentChunk.source_id == source_id,
            DocumentChunk.file_path == url
        ).first()
        if not chunk:
            chunk = db.query(DocumentChunk).filter(
                DocumentChunk.source_id == source_id,
                DocumentChunk.file_path.ilike(url)
            ).first()
        if chunk and chunk.metadata_json and "url" in chunk.metadata_json:
            url = chunk.metadata_json["url"]

    headers = {}
    auth = None
    if db_source.username and db_source.token:
        auth = (db_source.username, db_source.token)
    elif db_source.token:
        headers["Authorization"] = f"Bearer {db_source.token}"

    try:
        source_type = (db_source.type or "").lower()

        if source_type == "confluence":
            return _resolve_confluence(url, db_source.url, auth, headers, theme)
        elif source_type == "jira":
            return _resolve_jira(url, db_source.url, auth, headers, theme)
        else:
            raise HTTPException(status_code=400, detail="Diese Aktion wird nur für Confluence oder Jira unterstützt.")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Fehler beim Auflösen der URL: {str(e)}")


@router.get("/{source_id}/files")
def get_knowledge_source_files(
    source_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Gibt eine Liste aller eindeutigen Dateipfade der Wissensquelle zurück."""
    db_source = db.query(KnowledgeSource).filter(KnowledgeSource.id == source_id).first()
    if not db_source:
        raise HTTPException(status_code=404, detail="Wissensquelle nicht gefunden")
    assert_knowledge_source_visible(db_source, user, db)

    chunks = db.query(DocumentChunk.file_path).filter(DocumentChunk.source_id == source_id).distinct().all()
    # Eine Datei kann mehrere Chunks unter "<Dateipfad>#<suffix>" ablegen. Für den
    # Datei-Browser interessiert nur die zugrunde liegende Datei, nicht die einzelnen
    # Chunks — sonst erscheint eine Datei als Dutzende nicht öffenbare Einträge.
    files = sorted({c[0].split("#", 1)[0] for c in chunks if c[0]})
    return {"files": files}


@router.get("/{source_id}/content")
def get_knowledge_source_content(
    source_id: int,
    path: Optional[str] = None,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Liest den Inhalt einer Datei und gibt ihn als Text/HTML/SVG zurück.

    `path` wählt eine einzelne Datei innerhalb einer Mehrdatei-Quelle (Git,
    FolderWatch, WebDAV) aus; ohne `path` wird die in den Metadaten hinterlegte
    Einzeldatei einer "local"-Quelle verwendet.
    """
    db_source = db.query(KnowledgeSource).filter(KnowledgeSource.id == source_id).first()
    if not db_source:
        raise HTTPException(status_code=404, detail="Wissensquelle nicht gefunden")
    assert_knowledge_source_visible(db_source, user, db)

    path_name = (db_source.spaces or {}).get("path")
    if db_source.type == "Git" and path:
        # AP-3: Git-Quellen liegen als Worktree unter wt/ks_<id>, nicht mehr
        # flach unter REPOS_ROOT (Bare-Mirror + Worktree, siehe parser/git_utils.py).
        file_path = os.path.join(REPOS_ROOT, "wt", f"ks_{source_id}", path)
    elif path_name:
        # "local"-Quelle: einzelne Datei liegt unter dem in den Metadaten
        # hinterlegten Pfad, unabhängig vom angeklickten Dateibaum-Eintrag
        # (dessen Name für lokale Quellen nur der Anzeigename ist).
        file_path = os.path.join(UPLOADS_DIR, path_name)
    elif path:
        if path.startswith("watched/"):
            path = "/" + path
        file_path = path
    else:
        raise HTTPException(status_code=404, detail="Datei-Pfad nicht in Metadaten gefunden")

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Datei existiert nicht auf dem Server")

    try:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                content = "\n".join(p.extract_text() or "" for p in reader.pages)
                return {"content": content, "format": "text"}
            except Exception as e:
                # Fallback to plain text if the PDF is actually a text/mock file
                try:
                    with open(file_path, "r", errors="ignore") as f:
                        content = f.read()
                        if content.strip().startswith("##") or len(content) < 5000:
                            return {"content": content, "format": "text"}
                except Exception:
                    pass
                raise HTTPException(status_code=500, detail=f"Fehler beim Lesen der PDF-Datei: {str(e)}")

        elif ext in (".docx", ".doc"):
            try:
                import mammoth
                with open(file_path, "rb") as f:
                    result = mammoth.convert_to_html(f)
                return {"content": result.value, "format": "html"}
            except Exception as e:
                logger.warning(f"Mammoth HTML-Konvertierung fehlgeschlagen, Fallback auf Textmodus: {e}")
                import docx
                doc = docx.Document(file_path)
                parts = [p.text for p in doc.paragraphs]
                for table in doc.tables:
                    for row in table.rows:
                        parts.append(" | ".join(c.text for c in row.cells))
                return {"content": "\n".join(parts), "format": "text"}

        else:
            with open(file_path, "r", errors="ignore") as f:
                content = f.read()
            fmt = "markdown" if ext == ".md" else "text"
            return {"content": content, "format": fmt}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Fehler beim Lesen der Datei: {str(e)}")


@router.get("/{source_id}/raw")
def get_knowledge_source_raw(
    source_id: int,
    path: Optional[str] = None,
    download: bool = False,
    theme: Optional[str] = None,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Liefert eine Datei als Binary-Stream (für PDF-Viewer o.ä.)."""
    db_source = db.query(KnowledgeSource).filter(KnowledgeSource.id == source_id).first()
    if not db_source:
        raise HTTPException(status_code=404, detail="Wissensquelle nicht gefunden")
    assert_knowledge_source_visible(db_source, user, db)

    path_name = (db_source.spaces or {}).get("path")
    if path_name:
        file_path = os.path.join(UPLOADS_DIR, path_name)
    elif path:
        if path.startswith("watched/"):
            path = "/" + path
        file_path = path
    else:
        raise HTTPException(status_code=404, detail="Datei-Pfad nicht gefunden")

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Datei existiert nicht auf dem Server")

    ext = os.path.splitext(file_path)[1].lower()
    media_types = {".pdf": "application/pdf", ".png": "image/png", ".jpg": "image/jpeg",
                   ".jpeg": "image/jpeg", ".txt": "text/plain", ".md": "text/markdown"}
    media_type = media_types.get(ext, "application/octet-stream")

    if ext == ".pdf" and not download:
        try:
            with open(file_path, "rb") as f:
                header = f.read(4)
            if header != b"%PDF":
                # Mock PDF detected. Convert and serve as styled HTML.
                with open(file_path, "r", errors="ignore") as f:
                    content = f.read()
                import markdown
                html_body = markdown.markdown(content, extensions=['tables', 'fenced_code'])

                is_light = (theme == "light")
                bg_color = "#ffffff" if is_light else "#09090b"
                text_color = "#18181b" if is_light else "#d4d4d8"
                heading_color = "#09090b" if is_light else "#f4f4f5"
                border_color = "#e4e4e7" if is_light else "#27272a"
                th_bg = "#f4f4f5" if is_light else "#18181b"
                th_color = "#71717a" if is_light else "#a1a1aa"
                td_border = "#f4f4f5" if is_light else "#18181b"
                row_hover = "rgba(244, 244, 245, 0.5)" if is_light else "rgba(39, 39, 42, 0.2)"
                code_bg = "#f4f4f5" if is_light else "#18181b"
                code_color = "#e11d48" if is_light else "#f43f5e"
                scrollbar_thumb = "rgba(161, 161, 170, 0.4)" if is_light else "rgba(82, 82, 91, 0.5)"
                scrollbar_thumb_hover = "rgba(113, 113, 122, 0.65)" if is_light else "rgba(113, 113, 122, 0.75)"

                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <title>{os.path.basename(file_path)}</title>
                    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700&display=swap" rel="stylesheet">
                    <style>
                        body {{
                            font-family: 'Outfit', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
                            background-color: {bg_color};
                            color: {text_color};
                            padding: 2.5rem;
                            line-height: 1.6;
                            margin: 0;
                        }}
                        h1, h2, h3, h4 {{
                            color: {heading_color};
                            font-weight: 600;
                            margin-top: 1.5rem;
                            margin-bottom: 0.75rem;
                        }}
                        h1 {{
                            font-size: 1.75rem;
                            border-bottom: 1px solid {border_color};
                            padding-bottom: 0.5rem;
                            margin-top: 0;
                        }}
                        h2 {{ font-size: 1.4rem; }}
                        h3 {{ font-size: 1.15rem; }}
                        table {{
                            width: 100%;
                            border-collapse: collapse;
                            margin: 1.5rem 0;
                            font-size: 0.875rem;
                        }}
                        th {{
                            background-color: {th_bg};
                            color: {th_color};
                            font-weight: 500;
                            text-align: left;
                            padding: 0.75rem 1rem;
                            border-bottom: 2px solid {border_color};
                        }}
                        td {{
                            padding: 0.75rem 1rem;
                            border-bottom: 1px solid {td_border};
                        }}
                        tr:hover td {{
                            background-color: {row_hover};
                        }}
                        code {{
                            font-family: monospace;
                            background-color: {code_bg};
                            padding: 0.2rem 0.4rem;
                            border-radius: 0.25rem;
                            color: {code_color};
                        }}
                        p {{
                            margin-bottom: 1.25rem;
                        }}
                        ul, ol {{
                            padding-left: 1.5rem;
                            margin-bottom: 1.25rem;
                        }}
                        li {{
                            margin-bottom: 0.5rem;
                        }}
                        html {{
                            scrollbar-width: thin;
                            scrollbar-color: {scrollbar_thumb} transparent;
                        }}
                        ::-webkit-scrollbar {{
                            width: 10px;
                            height: 10px;
                        }}
                        ::-webkit-scrollbar-track {{
                            background: transparent;
                        }}
                        ::-webkit-scrollbar-thumb {{
                            background-color: {scrollbar_thumb};
                            border-radius: 9999px;
                            border: 2px solid transparent;
                            background-clip: padding-box;
                        }}
                        ::-webkit-scrollbar-thumb:hover {{
                            background-color: {scrollbar_thumb_hover};
                        }}
                        ::-webkit-scrollbar-corner {{
                            background: transparent;
                        }}
                    </style>
                </head>
                <body>
                    {html_body}
                </body>
                </html>
                """
                return HTMLResponse(content=html_content)
        except Exception as e:
            print("Failed to convert mock PDF to HTML:", e)

    if download:
        filename_to_use = os.path.basename(file_path) if path else db_source.name
        return FileResponse(file_path, media_type=media_type, filename=filename_to_use)
    return FileResponse(file_path, media_type=media_type, headers={"Content-Disposition": "inline"})


@router.post("/folder")
def create_folder_watch_source(
    source: FolderWatchCreate,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Registriert einen Ordner als Wissensquelle und startet den ersten Scan."""
    team_id = _resolve_team_id(source.project_id, db, user, source.team_id)
    _check_knowledge_source_cap(source.project_id, db)
    db_source = KnowledgeSource(
        name=source.name,
        type="FolderWatch",
        url=source.folder_path,
        project_id=source.project_id,
        sync_interval_minutes=_validate_sync_interval(source.sync_interval_minutes),
        team_id=team_id,
    )
    db.add(db_source)
    db.commit()
    db.refresh(db_source)
    send_tracked_task(db, db_source, "process_knowledge_source", [db_source.id], {"trace_id": get_trace_id()})
    return serialize_source(db_source)


@router.post("/git")
def create_git_source(
    source: GitSourceCreate,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Registriert ein Git-Repository als Wissensquelle und startet den ersten Scan."""
    team_id = _resolve_team_id(source.project_id, db, user, source.team_id)
    _check_knowledge_source_cap(source.project_id, db)
    db_source = KnowledgeSource(
        name=source.name,
        type="Git",
        url=source.url,
        username=source.username,
        token=source.token,
        project_id=source.project_id,
        # F-019: branch ist eine eigene Spalte (nicht mehr in spaces versteckt),
        # damit zwei Quellen auf dasselbe Repo mit verschiedenen Branches über
        # UNIQUE(project_id, url, branch) unterscheidbar bleiben. repo_fingerprint
        # setzt der Git-Konnektor beim ersten Sync (parser/connectors/git.py).
        branch=source.branch,
        spaces={"sparse_paths": source.sparse_paths},
        sync_interval_minutes=_validate_sync_interval(source.sync_interval_minutes),
        team_id=team_id,
    )
    db.add(db_source)
    db.commit()
    db.refresh(db_source)
    send_tracked_task(db, db_source, "process_knowledge_source", [db_source.id], {"trace_id": get_trace_id()})
    return serialize_source(db_source)



# Muss mit der `accept`-Liste in frontend/components/settings/tabs/SourcesSetupTab.tsx
# synchron bleiben — die UI begrenzt nur die Auswahl, hier wird sie durchgesetzt.
# O-044: an die vom Folder-/WebDAV-Connector unterstützten Formate angeglichen
# (parser/connectors/folder.py::SUPPORTED_EXTENSIONS) -- vorher konnte ein
# Kunde .docx nur über den Ordner-Watch-Connector einbinden, obwohl der lokale
# Upload-Pfad (parser/tasks/document.py) .docx längst extrahieren konnte.
_ALLOWED_UPLOAD_EXTENSIONS = {".pdf", ".md", ".txt", ".docx", ".doc"}


@router.post("/upload")
async def upload_local_document(
    file: UploadFile = File(...),
    name: str = Form(...),
    project_id: Optional[int] = Form(None),
    team_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    safe_filename = os.path.basename((file.filename or "upload").replace("\\", "/"))
    extension = os.path.splitext(safe_filename)[1].lower()
    if extension not in _ALLOWED_UPLOAD_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Dateityp '{extension or '(ohne Endung)'}' ist nicht erlaubt. "
                   f"Erlaubt: {', '.join(sorted(_ALLOWED_UPLOAD_EXTENSIONS))}",
        )

    resolved_team_id = _resolve_team_id(project_id, db, user, team_id)
    _check_knowledge_source_cap(project_id, db)
    db_source = KnowledgeSource(
        name=name,
        type="Local",
        project_id=project_id,
        team_id=resolved_team_id
    )
    db.add(db_source)
    db.commit()
    db.refresh(db_source)

    os.makedirs(UPLOADS_DIR, exist_ok=True)
    db_source.spaces = {"filename": safe_filename, "path": f"{db_source.id}_{safe_filename}"}
    db.commit()

    file_dest = os.path.join(UPLOADS_DIR, f"{db_source.id}_{safe_filename}")
    with open(file_dest, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    send_tracked_task(db, db_source, "process_local_document", [db_source.id, file_dest], {"trace_id": get_trace_id()})
    return serialize_source(db_source)


# ── Interne Render-Helfer für /resolve ───────────────────────────────────────

_DARK_STYLE = """<style>
body { font-family: ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif; line-height: 1.6; color: #d4d4d8; background-color: #09090b; margin: 0; padding: 24px; }
h1, h2, h3, h4, h5, h6 { color: #f4f4f5; font-weight: 700; margin-top: 1.5em; margin-bottom: 0.5em; }
a { color: #818cf8; text-decoration: none; } a:hover { text-decoration: underline; }
pre, code { background-color: #18181b; color: #f4f4f5; font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, monospace; padding: 0.2em 0.4em; border-radius: 4px; font-size: 85%; }
pre { padding: 1em; overflow-x: auto; }
img { max-width: 100%; height: auto; }
table { border-collapse: collapse; width: 100%; margin: 1.5em 0; }
th, td { border: 1px solid #27272a; padding: 8px 12px; text-align: left; }
th { background-color: #18181b; color: #f4f4f5; }
html { scrollbar-width: thin; scrollbar-color: rgba(92, 102, 117, 0.5) transparent; }
::-webkit-scrollbar { width: 10px; height: 10px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background-color: rgba(92, 102, 117, 0.5); border-radius: 9999px; border: 2px solid transparent; background-clip: padding-box; transition: background-color 0.15s; }
::-webkit-scrollbar-thumb:hover { background-color: rgba(92, 102, 117, 0.75); }
::-webkit-scrollbar-corner { background: transparent; }
</style>"""

_LIGHT_STYLE = """<style>
body { font-family: ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif; line-height: 1.6; color: #27272a; background-color: #ffffff; margin: 0; padding: 24px; }
h1, h2, h3, h4, h5, h6 { color: #09090b; font-weight: 700; margin-top: 1.5em; margin-bottom: 0.5em; }
a { color: #4f46e5; text-decoration: none; } a:hover { text-decoration: underline; }
pre, code { background-color: #f4f4f5; color: #09090b; font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, monospace; padding: 0.2em 0.4em; border-radius: 4px; font-size: 85%; }
pre { padding: 1em; overflow-x: auto; }
img { max-width: 100%; height: auto; }
table { border-collapse: collapse; width: 100%; margin: 1.5em 0; }
th, td { border: 1px solid #e4e4e7; padding: 8px 12px; text-align: left; }
th { background-color: #f4f4f5; color: #09090b; }
html { scrollbar-width: thin; scrollbar-color: rgba(161, 161, 170, 0.5) transparent; }
::-webkit-scrollbar { width: 10px; height: 10px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background-color: rgba(161, 161, 170, 0.5); border-radius: 9999px; border: 2px solid transparent; background-clip: padding-box; transition: background-color 0.15s; }
::-webkit-scrollbar-thumb:hover { background-color: rgba(161, 161, 170, 0.75); }
::-webkit-scrollbar-corner { background: transparent; }
</style>"""


def _resolve_confluence(url: str, base_url: str, auth, headers: dict, theme: str = "dark") -> dict:
    import urllib.parse as urlparse
    parsed = urlparse.urlparse(url)
    params = urlparse.parse_qs(parsed.query)

    page_id = None
    if "pageId" in params:
        page_id = params["pageId"][0]
    else:
        m = re.search(r'/pages/(\d+)', parsed.path)
        if m:
            page_id = m.group(1)
        else:
            for seg in parsed.path.strip("/").split("/"):
                if seg.isdigit() and len(seg) >= 5:
                    page_id = seg
                    break

    html_content = None
    with httpx.Client(timeout=15.0) as client:
        if page_id:
            for prefix in ["/wiki/rest/api/content", "/rest/api/content"]:
                try:
                    api_url = f"{base_url.rstrip('/')}{prefix}/{page_id}"
                    resp = client.get(api_url, auth=auth, headers=headers, params={"expand": "body.view"})
                    if resp.status_code == 200:
                        html_content = resp.json().get("body", {}).get("view", {}).get("value", "")
                        break
                except Exception:
                    continue
            if not html_content:
                raise HTTPException(status_code=404, detail="Confluence Seite konnte nicht abgerufen werden.")
        else:
            base_parsed = urlparse.urlparse(base_url)
            if (parsed.scheme, parsed.netloc) != (base_parsed.scheme, base_parsed.netloc):
                raise HTTPException(
                    status_code=400,
                    detail="URL liegt außerhalb der konfigurierten Confluence-Instanz.",
                )
            resp = client.get(url, auth=auth, headers=headers)
            resp.raise_for_status()
            html_content = resp.text

    clean = re.sub(r'<script\b[^<]*(?:(?!<\/script>)<[^<]*)*<\/script>', '', html_content)
    clean = re.sub(r'<a\b([^>]*)>', r'<a \1 target="_blank">', clean)
    style = _LIGHT_STYLE if theme == "light" else _DARK_STYLE
    return {"content": f"<html><head>{style}</head><body>{clean}</body></html>", "format": "html", "url": url}


def _resolve_jira(url: str, base_url: str, auth, headers: dict, theme: str = "dark") -> dict:
    import urllib.parse as urlparse
    parsed = urlparse.urlparse(url)
    params = urlparse.parse_qs(parsed.query)

    issue_key = None
    if "selectedIssue" in params:
        issue_key = params["selectedIssue"][0]
    else:
        m = re.search(r'/browse/([A-Z0-9]+-\d+)', parsed.path, re.IGNORECASE)
        if m:
            issue_key = m.group(1).upper()
        else:
            for seg in parsed.path.strip("/").split("/"):
                parts = seg.split("-")
                if len(parts) == 2 and parts[0].isalnum() and parts[1].isdigit():
                    issue_key = seg.upper()
                    break

    if not issue_key:
        raise HTTPException(status_code=400, detail="Kein Jira-Issue Key in der URL gefunden.")

    issue_data = None
    with httpx.Client(timeout=15.0) as client:
        for prefix in ["/rest/api/3/issue", "/rest/api/2/issue"]:
            try:
                resp = client.get(f"{base_url.rstrip('/')}{prefix}/{issue_key}", auth=auth, headers=headers)
                if resp.status_code == 200:
                    issue_data = resp.json()
                    break
            except Exception:
                continue

    if not issue_data:
        raise HTTPException(status_code=404, detail=f"Jira Ticket {issue_key} konnte nicht geladen werden.")

    fields = issue_data.get("fields", {})
    summary = fields.get("summary", "Kein Titel")
    status = fields.get("status", {}).get("name", "Unbekannt")
    priority = fields.get("priority", {}).get("name", "Mittel")
    assignee = fields.get("assignee", {}).get("displayName", "Nicht zugewiesen")

    desc_html = _adf_to_html(fields.get("description"))

    if theme == "light":
        body_bg = "#ffffff"
        body_color = "#27272a"
        badge_bg = "#f4f4f5"
        badge_color = "#09090b"
        badge_status_bg = "#e0e7ff"
        badge_status_color = "#3730a3"
        section_title_color = "#71717a"
        section_title_border = "#e4e4e7"
        pre_bg = "#f4f4f5"
        pre_color = "#09090b"
        header_border = "#e4e4e7"
        header_color = "#09090b"
        header_meta_color = "#71717a"
        comment_bg = "#f4f4f5"
        comment_border = "#e4e4e7"
        comment_meta_color = "#71717a"
        italic_color = "#71717a"
        scrollbar_thumb = "rgba(161, 161, 170, 0.5)"
        scrollbar_thumb_hover = "rgba(161, 161, 170, 0.75)"
    else:
        body_bg = "#09090b"
        body_color = "#d4d4d8"
        badge_bg = "#27272a"
        badge_color = "#f4f4f5"
        badge_status_bg = "#312e81"
        badge_status_color = "#c7d2fe"
        section_title_color = "#a1a1aa"
        section_title_border = "#18181b"
        pre_bg = "#18181b"
        pre_color = "#f4f4f5"
        header_border = "#27272a"
        header_color = "#f4f4f5"
        header_meta_color = "#a1a1aa"
        comment_bg = "#18181b"
        comment_border = "#27272a"
        comment_meta_color = "#a1a1aa"
        italic_color = "#71717a"
        scrollbar_thumb = "rgba(92, 102, 117, 0.5)"
        scrollbar_thumb_hover = "rgba(92, 102, 117, 0.75)"

    comments_html = ""
    for c in fields.get("comment", {}).get("comments", []):
        author = c.get("author", {}).get("displayName", "User")
        created = c.get("created", "")[:10]
        body_html = _adf_to_html(c.get("body")) if isinstance(c.get("body"), dict) else str(c.get("body", "")).replace("\n", "<br>")
        comments_html += (
            f'<div style="background-color:{comment_bg};border:1px solid {comment_border};border-radius:8px;padding:12px 16px;margin-bottom:12px;">'
            f'<div style="font-size:11px;font-weight:600;color:{comment_meta_color};margin-bottom:6px;">{author} &bull; {created}</div>'
            f'<div style="font-size:13px;">{body_html}</div></div>'
        )

    jira_style = f"""<style>
body {{ font-family: ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif; line-height: 1.6; color: {body_color}; background-color: {body_bg}; margin: 0; padding: 24px; }}
.badge {{ background-color: {badge_bg}; color: {badge_color}; padding: 2px 8px; border-radius: 4px; font-weight: 600; }}
.badge-status {{ background-color: {badge_status_bg}; color: {badge_status_color}; }}
.section-title {{ font-size: 14px; font-weight: 700; text-transform: uppercase; color: {section_title_color}; margin-top: 24px; margin-bottom: 8px; border-bottom: 1px solid {section_title_border}; padding-bottom: 4px; }}
pre, code {{ background-color: {pre_bg}; color: {pre_color}; font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, monospace; padding: 0.2em 0.4em; border-radius: 4px; font-size: 85%; }}
pre {{ padding: 1em; overflow-x: auto; }}
html {{ scrollbar-width: thin; scrollbar-color: {scrollbar_thumb} transparent; }}
::-webkit-scrollbar {{ width: 10px; height: 10px; }}
::-webkit-scrollbar-track {{ background: transparent; }}
::-webkit-scrollbar-thumb {{ background-color: {scrollbar_thumb}; border-radius: 9999px; border: 2px solid transparent; background-clip: padding-box; transition: background-color 0.15s; }}
::-webkit-scrollbar-thumb:hover {{ background-color: {scrollbar_thumb_hover}; }}
::-webkit-scrollbar-corner {{ background: transparent; }}
</style>"""

    num_comments = len(fields.get("comment", {}).get("comments", []))
    html = f"""<html><head>{jira_style}</head><body>
<div style="border-bottom:1px solid {header_border};padding-bottom:16px;margin-bottom:24px;">
  <h1 style="font-size:20px;font-weight:700;color:{header_color};margin:0 0 8px 0;">{issue_key}: {summary}</h1>
  <div style="display:flex;gap:16px;font-size:12px;color:{header_meta_color};">
    <span>Status: <span class="badge badge-status">{status}</span></span>
    <span>Priorität: <span class="badge">{priority}</span></span>
    <span>Zuständig: <span class="badge">{assignee}</span></span>
  </div>
</div>
<div class="section-title">Beschreibung</div>
<div style="font-size:14px;">{desc_html or f'<span style="color:{italic_color};font-style:italic;">Keine Beschreibung angegeben.</span>'}</div>
<div class="section-title">Kommentare ({num_comments})</div>
{comments_html or f'<span style="color:{italic_color};font-style:italic;">Keine Kommentare vorhanden.</span>'}
</body></html>"""

    return {"content": html, "format": "html", "url": url}


def _adf_to_html(node) -> str:
    """Konvertiert Atlassian Document Format (ADF) rekursiv in HTML."""
    if not node:
        return ""
    if node.get("type") == "text":
        text = node.get("text", "")
        for mark in node.get("marks", []):
            if mark.get("type") == "strong":
                text = f"<strong>{text}</strong>"
            elif mark.get("type") == "em":
                text = f"<em>{text}</em>"
        return text

    inner = "".join(_adf_to_html(c) for c in node.get("content", []))
    t = node.get("type")
    if t == "paragraph":
        return f"<p>{inner}</p>"
    elif t == "heading":
        lvl = node.get("attrs", {}).get("level", 3)
        return f"<h{lvl}>{inner}</h{lvl}>"
    elif t == "bulletList":
        return f"<ul>{inner}</ul>"
    elif t == "orderedList":
        return f"<ol>{inner}</ol>"
    elif t == "listItem":
        return f"<li>{inner}</li>"
    elif t == "codeBlock":
        return f"<pre><code>{inner}</code></pre>"
    elif t == "blockquote":
        return f"<blockquote>{inner}</blockquote>"
    return inner
